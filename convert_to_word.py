#!/usr/bin/env python3
"""
Comprehensive Markdown → Word (.docx) Converter
Converts all project docs, emails, and checklist to professional Word format.
Output: ~/ai-lending-company/word_docs/
"""
import os
import re
import html as html_mod
from pathlib import Path

import markdown
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = Path.home() / 'ai-lending-company'
OUTPUT = BASE / 'word_docs'
OUTPUT.mkdir(exist_ok=True)

# ── Styling Constants ──
PRIMARY = RGBColor(0x1A, 0x36, 0x5D)  # dark blue
ACCENT = RGBColor(0x05, 0x96, 0x69)  # green
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
RED = RGBColor(0xDC, 0x26, 0x26)

MD_EXTENSIONS = [
    'markdown.extensions.extra',
    'markdown.extensions.codehilite',
    'markdown.extensions.tables',
    'markdown.extensions.fenced_code',
    'markdown.extensions.smarty',
]


def add_formatted_text(paragraph, text, bold=False, italic=False, color=None,
                       size=None, font_name='Calibri'):
    """Add a run with formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(shading_elem)


def html_to_docx(doc, html_content):
    """Convert HTML to docx elements with proper formatting."""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    def process_children(parent_elem, container):
        """Process child elements recursively."""
        for child in parent_elem.children:
            if child.name is None:
                text = str(child).strip()
                if text:
                    p = container.add_paragraph()
                    add_formatted_text(p, text)
                continue

            if child.name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                level = int(child.name[1])
                p = container.add_paragraph()
                p.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
                p.paragraph_format.space_after = Pt(6)
                size = {1: 22, 2: 18, 3: 15, 4: 13, 5: 12, 6: 11}[level]
                color = PRIMARY if level <= 3 else BLACK
                add_formatted_text(p, child.get_text(), bold=True,
                                   color=color, size=size)

            elif child.name == 'p':
                p = container.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                for item in child.children:
                    _process_inline(item, p)

            elif child.name in ('ul', 'ol'):
                for li in child.find_all('li', recursive=False):
                    p = container.add_paragraph(style='List Bullet' if child.name == 'ul' else 'List Number')
                    p.paragraph_format.space_after = Pt(2)
                    for item in li.children:
                        _process_inline(item, p)

            elif child.name == 'pre':
                code_text = child.get_text()
                p = container.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
                # Add gray background shading
                shading = p._element.get_or_add_pPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'F5F5F5',
                })
                shading.append(shd)

            elif child.name == 'code':
                p = container.add_paragraph()
                run = p.add_run(child.get_text())
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xE8, 0x3E, 0x8C)

            elif child.name == 'blockquote':
                for inner in child.children:
                    if inner.name == 'p':
                        p = container.add_paragraph()
                        p.paragraph_format.left_indent = Cm(1)
                        p.paragraph_format.space_after = Pt(4)
                        add_formatted_text(p, '▎ ', color=ACCENT, size=11)
                        for item in inner.children:
                            _process_inline(item, p)

            elif child.name == 'hr':
                p = container.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run('─' * 60)
                run.font.color.rgb = GRAY
                run.font.size = Pt(8)

            elif child.name == 'table':
                rows = child.find_all('tr')
                if not rows:
                    continue
                cols = max(len(r.find_all(['th', 'td'])) for r in rows)
                table = container.add_table(rows=len(rows), cols=cols)
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for i, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for j, cell in enumerate(cells):
                        if j >= cols:
                            break
                        doc_cell = table.cell(i, j)
                        text = cell.get_text().strip()
                        doc_cell.text = ''
                        p = doc_cell.paragraphs[0]
                        is_header = cell.name == 'th'
                        add_formatted_text(p, text, bold=is_header,
                                           size=9 if not is_header else 10,
                                           color=WHITE if is_header else BLACK)
                        if is_header:
                            set_cell_shading(doc_cell, '1A365D')

            elif child.name == 'div' and child.get('class'):
                process_children(child, container)

    def _process_inline(item, paragraph):
        """Process inline elements within a paragraph."""
        if item.name is None:
            text = str(item)
            if text.strip():
                add_formatted_text(paragraph, text)
            return
        if item.name == 'strong':
            add_formatted_text(paragraph, item.get_text(), bold=True)
        elif item.name == 'em':
            add_formatted_text(paragraph, item.get_text(), italic=True)
        elif item.name == 'code':
            run = paragraph.add_run(item.get_text())
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xE8, 0x3E, 0x8C)
        elif item.name == 'a':
            text = item.get_text()
            href = item.get('href', '')
            add_formatted_text(paragraph, text, color=ACCENT)
            if href:
                add_formatted_text(paragraph, f' ({href})', color=GRAY, size=8)
        elif item.name in ('br',):
            paragraph.add_run('\n')
        elif item.name in ('ul', 'ol'):
            # Nested lists
            for li in item.find_all('li', recursive=False):
                add_formatted_text(paragraph, f'  • {li.get_text().strip()}', size=10)
                paragraph.add_run('\n')

    process_children(soup, doc)


def convert_md_to_docx(md_path, output_path=None):
    """Convert a single markdown file to Word."""
    if output_path is None:
        output_path = OUTPUT / md_path.with_suffix('.docx').name

    with open(md_path, 'r') as f:
        md_content = f.read()

    # Convert markdown to HTML
    html_content = markdown.markdown(md_content, extensions=MD_EXTENSIONS)

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Add a title from filename
    title = md_path.stem.replace('_', ' ').title()
    p = doc.add_heading(title, level=0)
    p.runs[0].font.color.rgb = PRIMARY
    p.runs[0].font.size = Pt(24)

    # Convert HTML to docx
    html_to_docx(doc, html_content)

    # Save
    doc.save(str(output_path))
    return output_path


def convert_email_to_docx(html_path, output_path=None):
    """Convert an HTML email template to Word."""
    if output_path is None:
        output_path = OUTPUT / f'email_{html_path.stem}.docx'

    with open(html_path, 'r') as f:
        html_content = f.read()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = f'Email Template: {html_path.stem.replace("_", " ").title()}'
    p = doc.add_heading(title, level=0)
    p.runs[0].font.color.rgb = ACCENT

    p = doc.add_paragraph()
    add_formatted_text(p, 'HTML Source Code:', bold=True, color=PRIMARY, size=12)

    # Wrap in code block
    for line in html_content.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(7)

    doc.save(str(output_path))
    return output_path


def create_combined_document():
    """Create a master document combining all key docs."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    add_formatted_text(p, 'SunCredit Lending Platform', bold=True, color=PRIMARY, size=32)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, 'Complete Business Documentation', color=ACCENT, size=18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    add_formatted_text(p, 'Generated May 2026', color=GRAY, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, f'Includes: Model Validation • Compliance • Operations • Legal • Pitch', color=GRAY, size=11)

    doc.add_page_break()

    # Table of Contents placeholder
    p = doc.add_heading('Table of Contents', level=1)
    p.runs[0].font.color.rgb = PRIMARY

    toc_entries = [
        ('READY_TO_LAUNCH.md', 'Master Launch Checklist'),
        ('ARCHITECTURE.md', 'System Architecture'),
        ('MODEL_VALIDATION_REPORT.md', 'Underwriting Model Validation'),
        ('MODEL_GOVERNANCE.md', 'Model Governance'),
        ('BANK_PARTNERSHIP_OUTREACH_EMAILS.md', 'Bank Partnership Emails'),
        ('AUTONOMOUS_OPERATION.md', 'Autonomous Operations'),
        ('INNOVATIVE_LENDING_PRODUCTS.md', 'Innovative Lending Products'),
        ('LOAN_CONTRACT.md', 'Loan Contract'),
        ('OPERATING_AGREEMENT.md', 'Operating Agreement'),
        ('VA_LLC_FORMATION_PACKET.md', 'Virginia LLC Formation'),
        ('COMPLIANCE_BUSINESS_PLAN.md', 'Compliance Business Plan'),
        ('COMPLIANCE_MANUAL_OUTLINE.md', 'Compliance Manual Outline'),
        ('PRIVATE_PLACEMENT_MEMO.md', 'Private Placement Memo'),
        ('PITCH_DECK_OUTLINE.md', 'Pitch Deck Outline'),
        ('SECURITY_AUDIT.md', 'Security Audit'),
        ('SECURITY_FIXES.md', 'Security Fixes'),
        ('DOMAIN_SETUP.md', 'Domain Setup'),
        ('PILOT_LOANS_GUIDE.md', 'Pilot Loans Guide'),
        ('TEXAS_LICENSING_RESEARCH.md', 'Texas Licensing Research'),
        ('SURETY_BOND_GUIDE.md', 'Surety Bond Guide'),
        ('KYC_VENDOR_EVAL.md', 'KYC Vendor Evaluation'),
        ('LENDER_PARTNERSHIP_OUTREACH.md', 'Lender Partnership Strategy'),
        ('PORTFOLIO_SUMMARY.md', 'Portfolio Summary'),
        ('Email Templates', '6 email templates (approved, welcome, payment, etc.)'),
    ]

    table = doc.add_table(rows=len(toc_entries), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (doc_name, desc) in enumerate(toc_entries):
        table.cell(i, 0).text = doc_name
        table.cell(i, 1).text = desc
        if i == 0:
            table.cell(i, 0).paragraphs[0].runs[0].bold = True if table.cell(i, 0).paragraphs[0].runs else False

    return doc


def main():
    print("=" * 60)
    print("  CONVERTING DOCUMENTS TO WORD FORMAT")
    print("=" * 60)

    # 1. Convert READY_TO_LAUNCH.md (master checklist)
    print("\n[1] Master Checklist...")
    path = BASE / 'READY_TO_LAUNCH.md'
    out = convert_md_to_docx(path)
    print(f"  ✓ {out.name}")

    # 2. Convert ARCHITECTURE.md
    path = BASE / 'ARCHITECTURE.md'
    if path.exists():
        out = convert_md_to_docx(path)
        print(f"  ✓ {out.name}")

    # 3. Convert RENDER_DEPLOY_STEPS.md
    path = BASE / 'RENDER_DEPLOY_STEPS.md'
    if path.exists():
        out = convert_md_to_docx(path)
        print(f"  ✓ {out.name}")

    # 4. Convert all launch/*.md
    print("\n[2] Launch Documents...")
    launch_dir = BASE / 'launch'
    md_files = sorted(launch_dir.glob('*.md'))
    for md_file in md_files:
        out = convert_md_to_docx(md_file)
        print(f"  ✓ {out.name}")

    # 5. Convert email templates
    print("\n[3] Email Templates...")
    email_dir = launch_dir / 'email_templates'
    if email_dir.exists():
        html_files = sorted(email_dir.glob('*.html'))
        for html_file in html_files:
            out = convert_email_to_docx(html_file)
            print(f"  ✓ {out.name}")

    # 6. Create combined document
    print("\n[4] Creating Combined Master Document...")
    combined = create_combined_document()

    # Append each document as a section
    for md_file in sorted(list(launch_dir.glob('*.md'))):
        with open(md_file) as f:
            html = markdown.markdown(f.read(), extensions=MD_EXTENSIONS)
        combined.add_page_break()
        p = combined.add_heading(md_file.stem.replace('_', ' ').title(), level=0)
        p.runs[0].font.color.rgb = PRIMARY
        html_to_docx(combined, html)

    combined_path = OUTPUT / 'SunCredit_Complete_Documentation.docx'
    combined.save(str(combined_path))
    print(f"  ✓ {combined_path.name}")

    # Summary
    docx_files = list(OUTPUT.glob('*.docx'))
    total_size = sum(f.stat().st_size for f in docx_files)
    print(f"\n{'='*60}")
    print(f"  COMPLETE: {len(docx_files)} Word documents created")
    print(f"  Total size: {total_size / 1024:.0f} KB")
    print(f"  Output: {OUTPUT}")
    print(f"  Combined doc: {combined_path.name} ({combined_path.stat().st_size / 1024:.0f} KB)")
    print(f"{'='*60}")


if __name__ == '__main__':
    main()
